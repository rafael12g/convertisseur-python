import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image
import os
from pathlib import Path

# Pour PDF
try:
    from PyPDF2 import PdfReader, PdfWriter
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Pour DOCX
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Pour Audio
try:
    from pydub import AudioSegment
    AUDIO_AVAILABLE = True
except ImportError:
    AUDIO_AVAILABLE = False

# Pour Vidéo
try:
    from moviepy.editor import VideoFileClip
    VIDEO_AVAILABLE = True
except ImportError:
    VIDEO_AVAILABLE = False


class ConvertisseurFichiersApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🔄 Convertisseur Universel Pro")
        self.root.geometry("700x650")
        self.root.configure(bg="#2c3e50")
        
        self.fichiers_selectionnes = []
        
        # Titre principal
        titre = tk.Label(root, text="🔄 CONVERTISSEUR UNIVERSEL", 
                        font=("Arial", 24, "bold"),
                        bg="#2c3e50", fg="white")
        titre.pack(pady=20)
        
        # Sous-titre
        sous_titre = tk.Label(root, text="Images • Documents • Audio • Vidéo", 
                            font=("Arial", 12),
                            bg="#2c3e50", fg="#95a5a6")
        sous_titre.pack(pady=5)
        
        # Frame principal (menu)
        self.frame_principal = tk.Frame(root, bg="#2c3e50")
        self.frame_principal.pack(pady=20)
        
        # Boutons de catégorie
        self.creer_boutons_categorie()
        
        # Canvas avec scrollbar pour frame de conversion
        self.canvas = tk.Canvas(root, bg="#34495e", highlightthickness=0)
        scrollbar = tk.Scrollbar(root, orient="vertical", command=self.canvas.yview)
        
        self.frame_conversion = tk.Frame(self.canvas, bg="#34495e", padx=20, pady=20)
        
        self.frame_conversion.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas.create_window((0, 0), window=self.frame_conversion, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar.set)
        
        self.scrollbar = scrollbar
        
        # Support de la molette
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        
    def _on_mousewheel(self, event):
        """Support de la molette de la souris"""
        if self.canvas.winfo_ismapped():
            self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
    
    def creer_boutons_categorie(self):
        """Crée les boutons pour choisir le type de conversion"""
        categories = [
            ("🖼️ Convertir Images", self.afficher_conversion_images, "#3498db", True),
            ("📄 Convertir Documents", self.afficher_conversion_documents, "#9b59b6", PDF_AVAILABLE or DOCX_AVAILABLE),
            ("🎵 Convertir Audio", self.afficher_conversion_audio, "#e67e22", AUDIO_AVAILABLE),
            ("🎬 Convertir Vidéo", self.afficher_conversion_video, "#e74c3c", VIDEO_AVAILABLE),
        ]
        
        for texte, commande, couleur, disponible in categories:
            if not disponible:
                texte += " 🔒"
                couleur = "#7f8c8d"
            
            btn = tk.Button(self.frame_principal, text=texte,
                          command=commande if disponible else lambda t=texte: self.module_manquant(t),
                          font=("Arial", 14, "bold"),
                          bg=couleur, fg="white",
                          width=28, height=2,
                          relief="raised", bd=3,
                          cursor="hand2")
            btn.pack(pady=8)
            
            if disponible:
                btn.bind("<Enter>", lambda e, b=btn, c=couleur: b.config(bg=self.darken_color(c)))
                btn.bind("<Leave>", lambda e, b=btn, c=couleur: b.config(bg=c))
    
    def darken_color(self, color):
        """Assombrit une couleur hexadécimale"""
        color = color.lstrip('#')
        r, g, b = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
        r, g, b = max(0, r-30), max(0, g-30), max(0, b-30)
        return f'#{r:02x}{g:02x}{b:02x}'
    
    def module_manquant(self, nom_fonction):
        """Affiche un message pour les modules manquants"""
        messages = {
            "📄 Convertir Documents 🔒": 
                "Pour utiliser cette fonctionnalité, installez:\n\n"
                "pip install PyPDF2 PyMuPDF python-docx",
            "🎵 Convertir Audio 🔒": 
                "Pour utiliser cette fonctionnalité, installez:\n\n"
                "pip install pydub\n\n"
                "Et installez FFmpeg sur votre système:\n"
                "https://ffmpeg.org/download.html",
            "🎬 Convertir Vidéo 🔒": 
                "Pour utiliser cette fonctionnalité, installez:\n\n"
                "pip install moviepy"
        }
        
        message = messages.get(nom_fonction, "Module non disponible")
        messagebox.showinfo("Installation requise", message)
    
    # ============= CONVERSION IMAGES =============
    
    def afficher_conversion_images(self):
        """Interface pour convertir des images"""
        self.frame_principal.pack_forget()
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        
        for widget in self.frame_conversion.winfo_children():
            widget.destroy()
        
        self.fichiers_selectionnes = []
        
        # Titre
        tk.Label(self.frame_conversion, text="🖼️ Conversion d'Images",
                font=("Arial", 20, "bold"),
                bg="#34495e", fg="white").pack(pady=15)
        
        # Formats supportés
        tk.Label(self.frame_conversion, 
                text="✓ Formats: JPG • PNG • BMP • GIF • WEBP • TIFF • ICO",
                font=("Arial", 10),
                bg="#34495e", fg="#95a5a6").pack(pady=5)
        
        # Bouton sélection
        btn_selectionner = tk.Button(self.frame_conversion, 
                                     text="📁 Sélectionner Image(s)",
                                     command=self.selectionner_images,
                                     font=("Arial", 13, "bold"),
                                     bg="#9b59b6", fg="white",
                                     width=28, height=2,
                                     cursor="hand2")
        btn_selectionner.pack(pady=15)
        
        # Liste fichiers
        self.label_fichiers = tk.Label(self.frame_conversion, 
                                       text="Aucun fichier sélectionné",
                                       font=("Arial", 10),
                                       bg="#34495e", fg="white",
                                       wraplength=550,
                                       justify="left")
        self.label_fichiers.pack(pady=10)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Format sortie
        tk.Label(self.frame_conversion, text="📤 Convertir vers:",
                font=("Arial", 13, "bold"),
                bg="#34495e", fg="white").pack(pady=8)
        
        formats_images = ["JPG", "PNG", "BMP", "GIF", "WEBP", "TIFF", "ICO"]
        self.combo_format_image = ttk.Combobox(self.frame_conversion, 
                                               values=formats_images,
                                               font=("Arial", 12),
                                               state="readonly",
                                               width=18)
        self.combo_format_image.pack(pady=5)
        self.combo_format_image.current(0)
        
        # Qualité
        frame_qualite = tk.Frame(self.frame_conversion, bg="#34495e")
        frame_qualite.pack(pady=12)
        
        tk.Label(frame_qualite, text="🎨 Qualité (pour JPG):",
                font=("Arial", 11),
                bg="#34495e", fg="white").pack(side="left", padx=5)
        
        self.scale_qualite = tk.Scale(frame_qualite, 
                                      from_=1, to=100,
                                      orient="horizontal",
                                      bg="#34495e", fg="white",
                                      highlightthickness=0,
                                      length=220)
        self.scale_qualite.set(95)
        self.scale_qualite.pack(side="left")
        
        # Redimensionnement
        self.var_redim = tk.BooleanVar()
        check_redim = tk.Checkbutton(self.frame_conversion,
                                    text="📏 Redimensionner les images",
                                    variable=self.var_redim,
                                    font=("Arial", 11, "bold"),
                                    bg="#34495e", fg="white",
                                    selectcolor="#2c3e50",
                                    activebackground="#34495e",
                                    activeforeground="white")
        check_redim.pack(pady=10)
        
        # Dimensions
        frame_dim = tk.Frame(self.frame_conversion, bg="#34495e")
        frame_dim.pack(pady=8)
        
        tk.Label(frame_dim, text="Largeur:",
                font=("Arial", 10),
                bg="#34495e", fg="white").pack(side="left", padx=5)
        self.entry_largeur = tk.Entry(frame_dim, width=10, font=("Arial", 11))
        self.entry_largeur.insert(0, "1920")
        self.entry_largeur.pack(side="left", padx=5)
        
        tk.Label(frame_dim, text="Hauteur:",
                font=("Arial", 10),
                bg="#34495e", fg="white").pack(side="left", padx=5)
        self.entry_hauteur = tk.Entry(frame_dim, width=10, font=("Arial", 11))
        self.entry_hauteur.insert(0, "1080")
        self.entry_hauteur.pack(side="left", padx=5)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=3, bg="#7f8c8d").pack(fill="x", pady=20)
        
        # BOUTON CONVERTIR
        btn_convertir = tk.Button(self.frame_conversion, 
                                 text="🔄 CONVERTIR MAINTENANT",
                                 command=self.convertir_images,
                                 font=("Arial", 16, "bold"),
                                 bg="#27ae60", fg="white",
                                 width=32, height=3,
                                 relief="raised",
                                 bd=5,
                                 cursor="hand2")
        btn_convertir.pack(pady=20)
        btn_convertir.bind("<Enter>", lambda e: btn_convertir.config(bg="#229954"))
        btn_convertir.bind("<Leave>", lambda e: btn_convertir.config(bg="#27ae60"))
        
        # Progression
        self.label_progression = tk.Label(self.frame_conversion, 
                                         text="",
                                         font=("Arial", 11, "bold"),
                                         bg="#34495e", fg="#f39c12")
        self.label_progression.pack(pady=10)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Bouton retour
        self.ajouter_bouton_retour()
        
        self.canvas.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def selectionner_images(self):
        """Sélectionne des images"""
        fichiers = filedialog.askopenfilenames(
            title="Sélectionner des images",
            filetypes=[
                ("Images", "*.jpg *.jpeg *.png *.bmp *.gif *.webp *.tiff *.ico"),
                ("Tous les fichiers", "*.*")
            ]
        )
        
        if fichiers:
            self.fichiers_selectionnes = list(fichiers)
            
            if len(fichiers) <= 10:
                noms = "\n".join([os.path.basename(f) for f in fichiers])
            else:
                premiers = "\n".join([os.path.basename(f) for f in fichiers[:10]])
                noms = f"{premiers}\n... et {len(fichiers)-10} autres fichiers"
            
            self.label_fichiers.config(
                text=f"✅ {len(fichiers)} fichier(s) sélectionné(s):\n\n{noms}",
                fg="#2ecc71"
            )
    
    def convertir_images(self):
        """Convertit les images"""
        if not self.fichiers_selectionnes:
            messagebox.showwarning("Attention", "Veuillez d'abord sélectionner des fichiers!")
            return
        
        format_sortie = self.combo_format_image.get().lower()
        qualite = self.scale_qualite.get()
        
        dossier_sortie = filedialog.askdirectory(title="Choisir le dossier de destination")
        if not dossier_sortie:
            return
        
        succes = 0
        echecs = 0
        
        for i, fichier in enumerate(self.fichiers_selectionnes):
            try:
                img = Image.open(fichier)
                
                # Redimensionner
                if self.var_redim.get():
                    try:
                        largeur = int(self.entry_largeur.get())
                        hauteur = int(self.entry_hauteur.get())
                        img = img.resize((largeur, hauteur), Image.Resampling.LANCZOS)
                    except ValueError:
                        messagebox.showerror("Erreur", "Dimensions invalides!")
                        return
                
                # Convertir en RGB pour JPG
                if format_sortie in ['jpg', 'jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'RGBA':
                        rgb_img.paste(img, mask=img.split()[-1])
                    else:
                        rgb_img.paste(img)
                    img = rgb_img
                
                # Nom de sortie
                nom_base = Path(fichier).stem
                nom_sortie = f"{nom_base}.{format_sortie}"
                chemin_sortie = os.path.join(dossier_sortie, nom_sortie)
                
                # Sauvegarder
                if format_sortie in ['jpeg', 'jpg']:
                    img.save(chemin_sortie, quality=qualite, optimize=True)
                else:
                    img.save(chemin_sortie)
                
                succes += 1
                
                progression = f"⏳ Conversion: {i+1}/{len(self.fichiers_selectionnes)}"
                self.label_progression.config(text=progression)
                self.root.update()
                
            except Exception as e:
                echecs += 1
                print(f"Erreur: {e}")
        
        message = f"✅ Conversion terminée!\n\n✓ Réussis: {succes}\n"
        if echecs > 0:
            message += f"✗ Échecs: {echecs}\n"
        message += f"\n📁 Fichiers dans:\n{dossier_sortie}"
        
        messagebox.showinfo("Terminé", message)
        self.label_progression.config(
            text=f"✅ Terminé! {succes}/{len(self.fichiers_selectionnes)}",
            fg="#2ecc71"
        )
    
    # ============= CONVERSION DOCUMENTS =============
    
    def afficher_conversion_documents(self):
        """Interface pour convertir des documents"""
        self.frame_principal.pack_forget()
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        
        for widget in self.frame_conversion.winfo_children():
            widget.destroy()
        
        self.fichiers_selectionnes = []
        
        # Titre
        tk.Label(self.frame_conversion, text="📄 Conversion de Documents",
                font=("Arial", 20, "bold"),
                bg="#34495e", fg="white").pack(pady=15)
        
        # Formats supportés
        formats_dispo = []
        if PDF_AVAILABLE:
            formats_dispo.append("PDF")
        if DOCX_AVAILABLE:
            formats_dispo.append("DOCX")
        formats_dispo.append("TXT")
        
        tk.Label(self.frame_conversion, 
                text=f"✓ Formats: {' • '.join(formats_dispo)}",
                font=("Arial", 10),
                bg="#34495e", fg="#95a5a6").pack(pady=5)
        
        # Bouton sélection
        btn_selectionner = tk.Button(self.frame_conversion, 
                                     text="📁 Sélectionner Document(s)",
                                     command=self.selectionner_documents,
                                     font=("Arial", 13, "bold"),
                                     bg="#9b59b6", fg="white",
                                     width=28, height=2,
                                     cursor="hand2")
        btn_selectionner.pack(pady=15)
        
        # Liste fichiers
        self.label_fichiers = tk.Label(self.frame_conversion, 
                                       text="Aucun fichier sélectionné",
                                       font=("Arial", 10),
                                       bg="#34495e", fg="white",
                                       wraplength=550,
                                       justify="left")
        self.label_fichiers.pack(pady=10)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Format sortie
        tk.Label(self.frame_conversion, text="📤 Convertir vers:",
                font=("Arial", 13, "bold"),
                bg="#34495e", fg="white").pack(pady=8)
        
        self.combo_format_doc = ttk.Combobox(self.frame_conversion, 
                                             values=formats_dispo,
                                             font=("Arial", 12),
                                             state="readonly",
                                             width=18)
        self.combo_format_doc.pack(pady=5)
        if formats_dispo:
            self.combo_format_doc.current(0)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=3, bg="#7f8c8d").pack(fill="x", pady=20)
        
        # BOUTON CONVERTIR
        btn_convertir = tk.Button(self.frame_conversion, 
                                 text="🔄 CONVERTIR MAINTENANT",
                                 command=self.convertir_documents,
                                 font=("Arial", 16, "bold"),
                                 bg="#27ae60", fg="white",
                                 width=32, height=3,
                                 relief="raised",
                                 bd=5,
                                 cursor="hand2")
        btn_convertir.pack(pady=20)
        btn_convertir.bind("<Enter>", lambda e: btn_convertir.config(bg="#229954"))
        btn_convertir.bind("<Leave>", lambda e: btn_convertir.config(bg="#27ae60"))
        
        # Progression
        self.label_progression = tk.Label(self.frame_conversion, 
                                         text="",
                                         font=("Arial", 11, "bold"),
                                         bg="#34495e", fg="#f39c12")
        self.label_progression.pack(pady=10)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Bouton retour
        self.ajouter_bouton_retour()
        
        self.canvas.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def selectionner_documents(self):
        """Sélectionne des documents"""
        types_fichiers = [("Documents", "*.pdf *.docx *.txt")]
        if PDF_AVAILABLE:
            types_fichiers.append(("PDF", "*.pdf"))
        if DOCX_AVAILABLE:
            types_fichiers.append(("Word", "*.docx"))
        types_fichiers.append(("Texte", "*.txt"))
        types_fichiers.append(("Tous les fichiers", "*.*"))
        
        fichiers = filedialog.askopenfilenames(
            title="Sélectionner des documents",
            filetypes=types_fichiers
        )
        
        if fichiers:
            self.fichiers_selectionnes = list(fichiers)
            
            if len(fichiers) <= 10:
                noms = "\n".join([os.path.basename(f) for f in fichiers])
            else:
                premiers = "\n".join([os.path.basename(f) for f in fichiers[:10]])
                noms = f"{premiers}\n... et {len(fichiers)-10} autres fichiers"
            
            self.label_fichiers.config(
                text=f"✅ {len(fichiers)} fichier(s) sélectionné(s):\n\n{noms}",
                fg="#2ecc71"
            )
    
    def convertir_documents(self):
        """Convertit les documents"""
        if not self.fichiers_selectionnes:
            messagebox.showwarning("Attention", "Veuillez d'abord sélectionner des fichiers!")
            return
        
        format_sortie = self.combo_format_doc.get().lower()
        
        dossier_sortie = filedialog.askdirectory(title="Choisir le dossier de destination")
        if not dossier_sortie:
            return
        
        succes = 0
        echecs = 0
        
        for i, fichier in enumerate(self.fichiers_selectionnes):
            try:
                extension = Path(fichier).suffix.lower()
                nom_base = Path(fichier).stem
                
                # PDF vers TXT
                if extension == '.pdf' and format_sortie == 'txt':
                    if PDF_AVAILABLE:
                        texte = ""
                        doc = fitz.open(fichier)
                        for page in doc:
                            texte += page.get_text()
                        doc.close()
                        
                        chemin_sortie = os.path.join(dossier_sortie, f"{nom_base}.txt")
                        with open(chemin_sortie, 'w', encoding='utf-8') as f:
                            f.write(texte)
                        succes += 1
                    else:
                        echecs += 1
                
                # DOCX vers TXT
                elif extension == '.docx' and format_sortie == 'txt':
                    if DOCX_AVAILABLE:
                        doc = Document(fichier)
                        texte = "\n".join([para.text for para in doc.paragraphs])
                        
                        chemin_sortie = os.path.join(dossier_sortie, f"{nom_base}.txt")
                        with open(chemin_sortie, 'w', encoding='utf-8') as f:
                            f.write(texte)
                        succes += 1
                    else:
                        echecs += 1
                
                # TXT vers DOCX
                elif extension == '.txt' and format_sortie == 'docx':
                    if DOCX_AVAILABLE:
                        with open(fichier, 'r', encoding='utf-8') as f:
                            contenu = f.read()
                        
                        doc = Document()
                        doc.add_paragraph(contenu)
                        
                        chemin_sortie = os.path.join(dossier_sortie, f"{nom_base}.docx")
                        doc.save(chemin_sortie)
                        succes += 1
                    else:
                        echecs += 1
                
                # PDF vers PDF (copie avec compression)
                elif extension == '.pdf' and format_sortie == 'pdf':
                    if PDF_AVAILABLE:
                        reader = PdfReader(fichier)
                        writer = PdfWriter()
                        
                        for page in reader.pages:
                            writer.add_page(page)
                        
                        chemin_sortie = os.path.join(dossier_sortie, f"{nom_base}_compressed.pdf")
                        with open(chemin_sortie, 'wb') as f:
                            writer.write(f)
                        succes += 1
                    else:
                        echecs += 1
                
                else:
                    echecs += 1
                
                progression = f"⏳ Conversion: {i+1}/{len(self.fichiers_selectionnes)}"
                self.label_progression.config(text=progression)
                self.root.update()
                
            except Exception as e:
                echecs += 1
                print(f"Erreur: {e}")
        
        message = f"✅ Conversion terminée!\n\n✓ Réussis: {succes}\n"
        if echecs > 0:
            message += f"✗ Échecs: {echecs}\n"
        message += f"\n📁 Fichiers dans:\n{dossier_sortie}"
        
        messagebox.showinfo("Terminé", message)
        self.label_progression.config(
            text=f"✅ Terminé! {succes}/{len(self.fichiers_selectionnes)}",
            fg="#2ecc71"
        )
    
    # ============= CONVERSION AUDIO =============
    
    def afficher_conversion_audio(self):
        """Interface pour convertir de l'audio"""
        self.frame_principal.pack_forget()
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        
        for widget in self.frame_conversion.winfo_children():
            widget.destroy()
        
        self.fichiers_selectionnes = []
        
        # Titre
        tk.Label(self.frame_conversion, text="🎵 Conversion Audio",
                font=("Arial", 20, "bold"),
                bg="#34495e", fg="white").pack(pady=15)
        
        # Formats supportés
        tk.Label(self.frame_conversion, 
                text="✓ Formats: MP3 • WAV • OGG • FLAC • AAC • M4A",
                font=("Arial", 10),
                bg="#34495e", fg="#95a5a6").pack(pady=5)
        
        # Bouton sélection
        btn_selectionner = tk.Button(self.frame_conversion, 
                                     text="📁 Sélectionner Audio",
                                     command=self.selectionner_audio,
                                     font=("Arial", 13, "bold"),
                                     bg="#9b59b6", fg="white",
                                     width=28, height=2,
                                     cursor="hand2")
        btn_selectionner.pack(pady=15)
        
        # Liste fichiers
        self.label_fichiers = tk.Label(self.frame_conversion, 
                                       text="Aucun fichier sélectionné",
                                       font=("Arial", 10),
                                       bg="#34495e", fg="white",
                                       wraplength=550,
                                       justify="left")
        self.label_fichiers.pack(pady=10)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Format sortie
        tk.Label(self.frame_conversion, text="📤 Convertir vers:",
                font=("Arial", 13, "bold"),
                bg="#34495e", fg="white").pack(pady=8)
        
        formats_audio = ["MP3", "WAV", "OGG", "FLAC", "AAC", "M4A"]
        self.combo_format_audio = ttk.Combobox(self.frame_conversion, 
                                               values=formats_audio,
                                               font=("Arial", 12),
                                               state="readonly",
                                               width=18)
        self.combo_format_audio.pack(pady=5)
        self.combo_format_audio.current(0)
        
        # Bitrate
        frame_bitrate = tk.Frame(self.frame_conversion, bg="#34495e")
        frame_bitrate.pack(pady=12)
        
        tk.Label(frame_bitrate, text="🎚️ Bitrate (kbps):",
                font=("Arial", 11),
                bg="#34495e", fg="white").pack(side="left", padx=5)
        
        bitrates = ["64", "128", "192", "256", "320"]
        self.combo_bitrate = ttk.Combobox(frame_bitrate, 
                                          values=bitrates,
                                          font=("Arial", 11),
                                          state="readonly",
                                          width=10)
        self.combo_bitrate.pack(side="left")
        self.combo_bitrate.current(3)  # 256 par défaut
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=3, bg="#7f8c8d").pack(fill="x", pady=20)
        
        # BOUTON CONVERTIR
        btn_convertir = tk.Button(self.frame_conversion, 
                                 text="🔄 CONVERTIR MAINTENANT",
                                 command=self.convertir_audio,
                                 font=("Arial", 16, "bold"),
                                 bg="#27ae60", fg="white",
                                 width=32, height=3,
                                 relief="raised",
                                 bd=5,
                                 cursor="hand2")
        btn_convertir.pack(pady=20)
        btn_convertir.bind("<Enter>", lambda e: btn_convertir.config(bg="#229954"))
        btn_convertir.bind("<Leave>", lambda e: btn_convertir.config(bg="#27ae60"))
        
        # Progression
        self.label_progression = tk.Label(self.frame_conversion, 
                                         text="",
                                         font=("Arial", 11, "bold"),
                                         bg="#34495e", fg="#f39c12")
        self.label_progression.pack(pady=10)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Bouton retour
        self.ajouter_bouton_retour()
        
        self.canvas.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def selectionner_audio(self):
        """Sélectionne des fichiers audio"""
        fichiers = filedialog.askopenfilenames(
            title="Sélectionner des fichiers audio",
            filetypes=[
                ("Audio", "*.mp3 *.wav *.ogg *.flac *.aac *.m4a *.wma"),
                ("Tous les fichiers", "*.*")
            ]
        )
        
        if fichiers:
            self.fichiers_selectionnes = list(fichiers)
            
            if len(fichiers) <= 10:
                noms = "\n".join([os.path.basename(f) for f in fichiers])
            else:
                premiers = "\n".join([os.path.basename(f) for f in fichiers[:10]])
                noms = f"{premiers}\n... et {len(fichiers)-10} autres fichiers"
            
            self.label_fichiers.config(
                text=f"✅ {len(fichiers)} fichier(s) sélectionné(s):\n\n{noms}",
                fg="#2ecc71"
            )
    
    def convertir_audio(self):
        """Convertit les fichiers audio"""
        if not AUDIO_AVAILABLE:
            messagebox.showerror("Erreur", "La bibliothèque pydub n'est pas installée!")
            return
        
        if not self.fichiers_selectionnes:
            messagebox.showwarning("Attention", "Veuillez d'abord sélectionner des fichiers!")
            return
        
        format_sortie = self.combo_format_audio.get().lower()
        bitrate = self.combo_bitrate.get() + "k"
        
        dossier_sortie = filedialog.askdirectory(title="Choisir le dossier de destination")
        if not dossier_sortie:
            return
        
        succes = 0
        echecs = 0
        
        for i, fichier in enumerate(self.fichiers_selectionnes):
            try:
                # Charger l'audio
                audio = AudioSegment.from_file(fichier)
                
                # Nom de sortie
                nom_base = Path(fichier).stem
                nom_sortie = f"{nom_base}.{format_sortie}"
                chemin_sortie = os.path.join(dossier_sortie, nom_sortie)
                
                # Exporter
                if format_sortie == "mp3":
                    audio.export(chemin_sortie, format="mp3", bitrate=bitrate)
                elif format_sortie == "wav":
                    audio.export(chemin_sortie, format="wav")
                elif format_sortie == "ogg":
                    audio.export(chemin_sortie, format="ogg", bitrate=bitrate)
                elif format_sortie == "flac":
                    audio.export(chemin_sortie, format="flac")
                elif format_sortie == "aac":
                    audio.export(chemin_sortie, format="adts", bitrate=bitrate)
                elif format_sortie == "m4a":
                    audio.export(chemin_sortie, format="mp4", bitrate=bitrate)
                else:
                    audio.export(chemin_sortie, format=format_sortie)
                
                succes += 1
                
                progression = f"⏳ Conversion: {i+1}/{len(self.fichiers_selectionnes)}"
                self.label_progression.config(text=progression)
                self.root.update()
                
            except Exception as e:
                echecs += 1
                print(f"Erreur: {e}")
        
        message = f"✅ Conversion terminée!\n\n✓ Réussis: {succes}\n"
        if echecs > 0:
            message += f"✗ Échecs: {echecs}\n"
        message += f"\n📁 Fichiers dans:\n{dossier_sortie}"
        
        messagebox.showinfo("Terminé", message)
        self.label_progression.config(
            text=f"✅ Terminé! {succes}/{len(self.fichiers_selectionnes)}",
            fg="#2ecc71"
        )
    
    # ============= CONVERSION VIDÉO =============
    
    def afficher_conversion_video(self):
        """Interface pour convertir des vidéos"""
        self.frame_principal.pack_forget()
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        
        for widget in self.frame_conversion.winfo_children():
            widget.destroy()
        
        self.fichiers_selectionnes = []
        
        # Titre
        tk.Label(self.frame_conversion, text="🎬 Conversion Vidéo",
                font=("Arial", 20, "bold"),
                bg="#34495e", fg="white").pack(pady=15)
        
        # Formats supportés
        tk.Label(self.frame_conversion, 
                text="✓ Formats: MP4 • AVI • MKV • MOV • WEBM • FLV",
                font=("Arial", 10),
                bg="#34495e", fg="#95a5a6").pack(pady=5)
        
        # Bouton sélection
        btn_selectionner = tk.Button(self.frame_conversion, 
                                     text="📁 Sélectionner Vidéo(s)",
                                     command=self.selectionner_video,
                                     font=("Arial", 13, "bold"),
                                     bg="#9b59b6", fg="white",
                                     width=28, height=2,
                                     cursor="hand2")
        btn_selectionner.pack(pady=15)
        
        # Liste fichiers
        self.label_fichiers = tk.Label(self.frame_conversion, 
                                       text="Aucun fichier sélectionné",
                                       font=("Arial", 10),
                                       bg="#34495e", fg="white",
                                       wraplength=550,
                                       justify="left")
        self.label_fichiers.pack(pady=10)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Format sortie
        tk.Label(self.frame_conversion, text="📤 Convertir vers:",
                font=("Arial", 13, "bold"),
                bg="#34495e", fg="white").pack(pady=8)
        
        formats_video = ["MP4", "AVI", "MKV", "MOV", "WEBM", "FLV"]
        self.combo_format_video = ttk.Combobox(self.frame_conversion, 
                                               values=formats_video,
                                               font=("Arial", 12),
                                               state="readonly",
                                               width=18)
        self.combo_format_video.pack(pady=5)
        self.combo_format_video.current(0)
        
        # Qualité/Codec
        frame_codec = tk.Frame(self.frame_conversion, bg="#34495e")
        frame_codec.pack(pady=12)
        
        tk.Label(frame_codec, text="🎥 Codec:",
                font=("Arial", 11),
                bg="#34495e", fg="white").pack(side="left", padx=5)
        
        codecs = ["libx264 (H.264)", "libx265 (H.265/HEVC)", "mpeg4", "copy"]
        self.combo_codec = ttk.Combobox(frame_codec, 
                                        values=codecs,
                                        font=("Arial", 11),
                                        state="readonly",
                                        width=20)
        self.combo_codec.pack(side="left")
        self.combo_codec.current(0)
        
        # Résolution
        tk.Label(self.frame_conversion, text="📐 Résolution:",
                font=("Arial", 11, "bold"),
                bg="#34495e", fg="white").pack(pady=8)
        
        resolutions = ["Original", "1920x1080 (1080p)", "1280x720 (720p)", 
                      "854x480 (480p)", "640x360 (360p)"]
        self.combo_resolution = ttk.Combobox(self.frame_conversion, 
                                             values=resolutions,
                                             font=("Arial", 11),
                                             state="readonly",
                                             width=25)
        self.combo_resolution.pack(pady=5)
        self.combo_resolution.current(0)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=3, bg="#7f8c8d").pack(fill="x", pady=20)
        
        # BOUTON CONVERTIR
        btn_convertir = tk.Button(self.frame_conversion, 
                                 text="🔄 CONVERTIR MAINTENANT",
                                 command=self.convertir_video,
                                 font=("Arial", 16, "bold"),
                                 bg="#27ae60", fg="white",
                                 width=32, height=3,
                                 relief="raised",
                                 bd=5,
                                 cursor="hand2")
        btn_convertir.pack(pady=20)
        btn_convertir.bind("<Enter>", lambda e: btn_convertir.config(bg="#229954"))
        btn_convertir.bind("<Leave>", lambda e: btn_convertir.config(bg="#27ae60"))
        
        # Progression
        self.label_progression = tk.Label(self.frame_conversion, 
                                         text="",
                                         font=("Arial", 11, "bold"),
                                         bg="#34495e", fg="#f39c12")
        self.label_progression.pack(pady=10)
        
        # Avertissement
        tk.Label(self.frame_conversion, 
                text="⚠️ La conversion vidéo peut prendre du temps",
                font=("Arial", 9, "italic"),
                bg="#34495e", fg="#e67e22").pack(pady=5)
        
        # Séparateur
        tk.Frame(self.frame_conversion, height=2, bg="#7f8c8d").pack(fill="x", pady=15)
        
        # Bouton retour
        self.ajouter_bouton_retour()
        
        self.canvas.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def selectionner_video(self):
        """Sélectionne des vidéos"""
        fichiers = filedialog.askopenfilenames(
            title="Sélectionner des vidéos",
            filetypes=[
                ("Vidéo", "*.mp4 *.avi *.mkv *.mov *.webm *.flv *.wmv *.m4v"),
                ("Tous les fichiers", "*.*")
            ]
        )
        
        if fichiers:
            self.fichiers_selectionnes = list(fichiers)
            
            if len(fichiers) <= 10:
                noms = "\n".join([os.path.basename(f) for f in fichiers])
            else:
                premiers = "\n".join([os.path.basename(f) for f in fichiers[:5]])
                noms = f"{premiers}\n... et {len(fichiers)-5} autres fichiers"
            
            self.label_fichiers.config(
                text=f"✅ {len(fichiers)} fichier(s) sélectionné(s):\n\n{noms}",
                fg="#2ecc71"
            )
    
    def convertir_video(self):
        """Convertit les vidéos"""
        if not VIDEO_AVAILABLE:
            messagebox.showerror("Erreur", "La bibliothèque moviepy n'est pas installée!")
            return
        
        if not self.fichiers_selectionnes:
            messagebox.showwarning("Attention", "Veuillez d'abord sélectionner des fichiers!")
            return
        
        format_sortie = self.combo_format_video.get().lower()
        codec_text = self.combo_codec.get()
        codec = codec_text.split()[0]  # Extraire juste le nom du codec
        
        resolution_text = self.combo_resolution.get()
        
        dossier_sortie = filedialog.askdirectory(title="Choisir le dossier de destination")
        if not dossier_sortie:
            return
        
        succes = 0
        echecs = 0
        
        for i, fichier in enumerate(self.fichiers_selectionnes):
            try:
                # Charger la vidéo
                video = VideoFileClip(fichier)
                
                # Redimensionner si demandé
                if resolution_text != "Original":
                    resolution = resolution_text.split()[0]  # "1920x1080"
                    largeur, hauteur = map(int, resolution.split('x'))
                    video = video.resize((largeur, hauteur))
                
                # Nom de sortie
                nom_base = Path(fichier).stem
                nom_sortie = f"{nom_base}.{format_sortie}"
                chemin_sortie = os.path.join(dossier_sortie, nom_sortie)
                
                # Exporter
                if codec == "copy":
                    video.write_videofile(chemin_sortie, codec="copy", audio_codec="copy")
                else:
                    video.write_videofile(chemin_sortie, codec=codec)
                
                video.close()
                succes += 1
                
                progression = f"⏳ Conversion: {i+1}/{len(self.fichiers_selectionnes)}"
                self.label_progression.config(text=progression)
                self.root.update()
                
            except Exception as e:
                echecs += 1
                print(f"Erreur: {e}")
        
        message = f"✅ Conversion terminée!\n\n✓ Réussis: {succes}\n"
        if echecs > 0:
            message += f"✗ Échecs: {echecs}\n"
        message += f"\n📁 Fichiers dans:\n{dossier_sortie}"
        
        messagebox.showinfo("Terminé", message)
        self.label_progression.config(
            text=f"✅ Terminé! {succes}/{len(self.fichiers_selectionnes)}",
            fg="#2ecc71"
        )
    
    # ============= UTILITAIRES =============
    
    def ajouter_bouton_retour(self):
        """Ajoute un bouton retour au menu"""
        btn_retour = tk.Button(self.frame_conversion, 
                              text="← Retour au menu principal",
                              command=self.retour_menu,
                              font=("Arial", 12, "bold"),
                              bg="#e74c3c", fg="white",
                              width=28, height=2,
                              cursor="hand2")
        btn_retour.pack(pady=20)
        btn_retour.bind("<Enter>", lambda e: btn_retour.config(bg="#c0392b"))
        btn_retour.bind("<Leave>", lambda e: btn_retour.config(bg="#e74c3c"))
    
    def retour_menu(self):
        """Retourne au menu principal"""
        self.canvas.pack_forget()
        self.scrollbar.pack_forget()
        self.frame_principal.pack(pady=20)
        self.fichiers_selectionnes = []


# ============= LANCEMENT =============

if __name__ == "__main__":
    root = tk.Tk()
    app = ConvertisseurFichiersApp(root)
    
    # Message de bienvenue avec statut des modules
    modules_status = []
    if not PDF_AVAILABLE:
        modules_status.append("📄 Documents: Installer PyPDF2 et PyMuPDF")
    if not DOCX_AVAILABLE:
        modules_status.append("📝 DOCX: Installer python-docx")
    if not AUDIO_AVAILABLE:
        modules_status.append("🎵 Audio: Installer pydub + FFmpeg")
    if not VIDEO_AVAILABLE:
        modules_status.append("🎬 Vidéo: Installer moviepy")
    
    if modules_status:
        print("\n" + "="*50)
        print("⚠️  MODULES OPTIONNELS MANQUANTS:")
        print("="*50)
        for msg in modules_status:
            print(f"  • {msg}")
        print("\nPour installer:")
        print("  pip install PyPDF2 PyMuPDF python-docx pydub moviepy")
        print("="*50 + "\n")
    
    root.mainloop()
